from pathlib import Path
from copy import deepcopy
import hashlib
import os
import tempfile
from zipfile import ZipFile

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BLUEPRINT_ROOT = Path(__file__).resolve().parents[1]
SOURCE = BLUEPRINT_ROOT / "3-业务蓝图-终态/SCOS_中粮太仓项目研发模块蓝图方案v0.1.docx"
OUTPUT = BLUEPRINT_ROOT / "3-业务蓝图-终态/SCOS_中粮太仓项目研发模块蓝图方案v0.2_业务蓝图初稿.docx"
ARCH_IMG = Path("/tmp/scos_blueprint_v02/pdfs/SCOS整体架构.png")
EXPECTED_SHA = "466d9be41e5821c8b75c6a92071e8dacde04b15dac7afe0985a45d230e9d8cfe"

TOTAL_DXA = 8280
HEADER_FILL = "D9E2F3"
SUB_FILL = "EAF0F7"
NOTE_FILL = "FFF2CC"
INK = RGBColor(31, 78, 121)

PRESERVE_TEMPLATE_PARTS = {
    "word/styles.xml",
    "word/numbering.xml",
    "word/header1.xml",
    "word/header2.xml",
    "word/_rels/header1.xml.rels",
    "word/_rels/header2.xml.rels",
    "word/footer1.xml",
    "word/footer2.xml",
    "word/footer3.xml",
    "word/footer4.xml",
}


def set_east_asia(run, name="宋体", size=10.5, bold=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=85, start=100, bottom=85, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def set_table_geometry(table, widths):
    assert sum(widths) == TOTAL_DXA
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TOTAL_DXA)); tbl_w.set(qn("w:type"), "dxa")
    ind = tbl_pr.find(qn("w:tblInd"))
    if ind is None:
        ind = OxmlElement("w:tblInd"); tbl_pr.append(ind)
    ind.set(qn("w:w"), "0"); ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(w)); grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW"); tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[i])); tc_w.set(qn("w:type"), "dxa")


def move_before(element, marker):
    marker.addprevious(element)


def add_p(doc, marker, text="", style="Body Text", bold_lead=None, align=None, after=5):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead); set_east_asia(r, bold=True)
        r = p.add_run(text[len(bold_lead):]); set_east_asia(r)
    else:
        r = p.add_run(text); set_east_asia(r)
    move_before(p._p, marker)
    return p


def add_heading(doc, marker, text, level, page_break=False):
    p = doc.add_paragraph(style=f"Heading {level}")
    if page_break:
        p.paragraph_format.page_break_before = True
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_east_asia(r, name="黑体", size={1:16,2:15,3:14}[level], bold=True, color=INK)
    move_before(p._p, marker)
    return p


def add_table(doc, marker, headers, rows, widths, font=8.5, header_fill=HEADER_FILL):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]; repeat_header(hdr)
    for i, h in enumerate(headers):
        hdr.cells[i].text = str(h)
        shade(hdr.cells[i], header_fill)
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cells[i].text = str(value)
    set_table_geometry(table, widths)
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(1.5)
                p.paragraph_format.line_spacing = 1.05
                for r in p.runs:
                    set_east_asia(r, size=font, bold=(ri == 0))
            if ri > 0 and ri % 2 == 0:
                shade(cell, "F8FAFC")
    move_before(table._tbl, marker)
    return table


def add_note(doc, marker, text, label="待确认"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0); shade(cell, NOTE_FILL); set_cell_margins(cell, 110, 130, 110, 130)
    p = cell.paragraphs[0]
    r = p.add_run(f"【{label}】"); set_east_asia(r, bold=True)
    r = p.add_run(text); set_east_asia(r)
    set_table_geometry(table, [TOTAL_DXA])
    move_before(table._tbl, marker)
    return table


def add_picture(doc, marker, path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(5.65))
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", caption)
    p.paragraph_format.keep_with_next = True
    move_before(p._p, marker)
    c = doc.add_paragraph(style="Body Text")
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(6)
    r = c.add_run(caption); set_east_asia(r, size=9)
    move_before(c._p, marker)


def find_body_para(doc, text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p._p
    raise ValueError(f"paragraph not found: {text}")


def remove_between(doc, start_text, end_text):
    start = find_body_para(doc, start_text)
    end = find_body_para(doc, end_text)
    body = doc._element.body
    elements = list(body.iterchildren())
    a, b = elements.index(start), elements.index(end)
    for element in elements[a:b]:
        body.remove(element)
    return end


def remove_from_to_sectpr(doc, start_element):
    body = doc._element.body
    elements = list(body.iterchildren())
    start = elements.index(start_element)
    sect_pr = body.sectPr
    for element in elements[start:]:
        if element is not sect_pr:
            body.remove(element)
    return sect_pr


def replace_para_text(p, text):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    r = p.add_run(text); set_east_asia(r, size=10)


def update_visible_toc(doc):
    toc = {
        24: "二、 SCOS业务蓝图交付物\t7",
        25: "2.1 SCOS业务流程总览\t7",
        40: "2.4 SCOS角色权限矩阵\t25",
        43: "2.5 SCOS业务规则清单\t32",
        44: "三、 SCOS智能化应用场景说明书\t页码待更新",
        45: "3.1 智能化场景总体说明\t页码待更新",
        46: "3.2 供应链驾驶舱态势监测与预警联动\t页码待更新",
        47: "3.3 供应链沙盘推演\t页码待更新",
        48: "3.4 采购、库存、排程、成本与追溯场景\t页码待更新",
        49: "四、 SCOS技术蓝图交付物\t页码待更新",
        50: "4.1 SCOS总体技术架构\t页码待更新",
        51: "4.2 SCOS数据流架构\t页码待更新",
        52: "4.3 SCOS数据域与核心实体\t页码待更新",
        53: "4.3.1 主数据域\t页码待更新",
        54: "4.3.2 需求、计划与生产数据域\t页码待更新",
        55: "4.3.3 库存、采购与供应数据域\t页码待更新",
        56: "4.3.4 财务、成本与追溯数据域\t页码待更新",
        57: "4.4 SCOS模型与计算服务\t页码待更新",
        58: "4.4.1 采购优化服务\t页码待更新",
        59: "4.4.2 成品库存优化服务\t页码待更新",
        60: "4.4.3 精益排程服务\t页码待更新",
        61: "4.4.4 驾驶舱与沙盘编排服务\t页码待更新",
        62: "4.5 数据发布、版本与审计规范\t页码待更新",
        63: "4.6 安全、运维与非功能\t页码待更新",
        64: "4.7 技术蓝图待确认项\t页码待更新",
        65: "",
        66: "",
        67: "",
        68: "",
        69: "",
        70: "五、 SCOS集成接口规范\t页码待更新",
        71: "六、 SCOS业务场景-智能化覆盖率映射表\t页码待更新",
        75: "七、 SCOS非功能性需求响应表\t页码待更新",
        76: "八、 SCOS蓝图交付物清单\t页码待更新",
    }
    for idx, text in toc.items():
        replace_para_text(doc.paragraphs[idx], text)


def add_scenario_template(doc, marker, title, value, roles, trigger, inputs, logic, outputs, rules, acceptance, pending=None):
    add_heading(doc, marker, title, 2)
    add_p(doc, marker, value)
    add_table(doc, marker, ["项目", "场景说明"], [
        ("使用角色", roles),
        ("触发方式", trigger),
        ("输入数据", inputs),
        ("核心处理", logic),
        ("输出及去向", outputs),
        ("业务规则", rules),
        ("验收要点", acceptance),
    ], [1500, 6780], font=8.6)
    if pending:
        add_note(doc, marker, pending)


def write_chapter_3(doc, marker):
    add_heading(doc, marker, "SCOS智能化应用场景说明书", 1, page_break=True)
    add_heading(doc, marker, "智能化场景总体说明", 2)
    add_p(doc, marker, "SCOS定位为工厂供应链智慧大脑，围绕“产销-库存-采购-生产-成本-追溯”形成数据驱动的全局优化与协同闭环。系统从统一数据服务层和实时数据平台消费业务数据，通过规则计算、优化模型和情景模拟形成可解释的决策建议，并经业务人员确认、审批和发布后进入执行系统。")
    add_p(doc, marker, "本章所述智能化不以“自动替代人员决策”为目标，而是将人工经验中可固化的口径、约束和评判标准沉淀为系统能力。每次计算均应保留输入数据版本、模型参数、结果方案、人工调整与发布记录，以满足复核、审计和追溯要求。")
    add_table(doc, marker, ["类别", "智能化场景", "主要输出", "当前完整度"], [
        ("驾驶舱", "全局态势监测、健康度评估、预警钻取与处置闭环", "核心KPI、专题视图、分级预警、态势报告", "较完整"),
        ("沙盘", "需求波动、供应延迟和原料价格变化推演", "当前方案与模拟方案影响对比", "方案级完整"),
        ("采购", "待采购需求净算、供应商组合优化、锁价和到货建议", "推荐/备选方案、建议下单与到货日期", "较完整"),
        ("库存", "成品安全库存、目标库存和月-周-日生产补货协同", "SS/s/S、补货量、风险日、期望入库日期", "成品范围内较完整"),
        ("排程", "4-6周粗排与24-72小时精排、异常快速重排", "原辅包需求、车间日计划、交付风险和重排方案", "总体方案可写，约束待补"),
        ("成本/追溯", "成本归集分摊、多维盈利分析、批次正反向追溯", "成本差异、客户/产品贡献、批次链路和质量联动", "总体方案可写，口径待补"),
    ], [1050, 3000, 2700, 1530], font=8.2)

    add_scenario_template(doc, marker,
        "供应链驾驶舱态势监测与预警联动",
        "本场景面向管理层和执行层提供“产销-储运-环保-资金”一站式态势视图，使管理者能够在同一口径下识别异常、追溯原因并跟踪处置结果。驾驶舱作为消费层和展示层，不自行绕过统一数据架构连接生产数据库。",
        "公司管理层、供应链计划员、生产调度员、采购专员、仓储和物流管理人员。",
        "系统按实时/定时频率自动刷新；关键指标越界、业务事件发生或用户主动查询时触发。",
        "统一数据服务层的主数据和交易数据；实时数据平台的产量、库存、储罐液位、产线状态和车辆位置；SAP财务、资金与成本数据；各SCOS模块计算结果和预警事件。",
        "指标计算引擎按统一口径计算KPI，与配置阈值比较后形成红黄绿健康度，越界时生成分级预警。用户可按“KPI-专题-明细-源单据/工单”链路钻取，并记录责任人、处置动作和关闭结果。",
        "今日产销达成率、成品库存周转天数、订单准时交付率、供应链总运营成本、现金循环周期等核心KPI；产销、储运、环保、资金四类专题视图；一/二级预警、根因链路、处置记录和态势报告。",
        "指标不得以临时手工录入替代系统计算；一级预警用于产线故障停机、关键原料断供、重大质量事故等紧急场景，二级预警用于库存低于阈值、交付延误风险、能耗超标等警告场景。",
        "主要视图首次加载不超过3秒，数据钻取响应不超过2秒；产量、发货量、库存、车辆和产线状态等关键运营指标延迟原则上不超过5秒；成本和资金类指标更新频率不超过1小时。",
        "需由业务部门确认最终KPI字典、红黄绿阈值、预警处置工单承接系统及移动端推送范围。")

    add_scenario_template(doc, marker,
        "供应链沙盘推演",
        "沙盘推演用于在不修改正式业务数据的前提下，量化需求、供应和成本变化对生产、库存、采购、交付和资金的连锁影响，帮助管理层在多个可行方案之间做出选择。",
        "公司管理层、供应链负责人、生产和采购业务负责人。",
        "用户主动进入沙盘并选择基线方案、设置参数后触发；也可从驾驶舱预警页携带当前业务快照进入。",
        "与驾驶舱相同时点的数据快照、已发布计划方案、当前库存和采购执行状态；用户输入的大客户订单增减比例、供应商延迟天数、原料价格变化比例等情景参数。",
        "沙盘编排服务复制基线快照，根据情景类型调用排程、采购、库存或成本模型，将返回结果重新汇总为全局KPI，并与当前基线对比。",
        "需求变化对产能、原料、库存和交付日期的影响；供应延迟下的断供风险、备用供应商或排程调整方案；原料价格变化下的采购与总运营成本影响。",
        "沙盘内的参数和结果不得直接改写正式计划；用户采纳模拟方案后，系统将其转换为待审批的调整建议，按相应模块的审批发布流程处理。",
        "同一数据快照和参数集应能重复得到一致结果；当被调模型超时、不可行或数据不完整时，系统应明确返回原因，不得输出误导性结论。",
        "需在各模型完成后确认沙盘跨模型调用顺序、超时标准、参数允许范围和方案采纳权限。")

    add_scenario_template(doc, marker,
        "原料采购决策优化",
        "采购优化场景解决“需要采购多少、向哪些供应商采购、各采购多少、何时下单和到货”等决策问题。系统以已发布的原材料需求计划为需求基线，通过净算已有采购承诺、比较供应条件形成可人工调整的采购建议。",
        "采购专员、采购负责人、财务人员、供应链计划员。",
        "每月根据新发布的原材料需求计划形成月度方案；每周检查需求和采购执行变化；计划版本变更或业务人员主动发起时重新计算未执行部分。",
        "已发布的原材料分期需求量和需求日期；已生效未完成采购订单及在途/到货状态；供应商报价、最大可供量、运费、账期、备货/运输/质检提前期；当前期货价格和资金成本率。",
        "按物料和需求周期扣减在需求日前可到货的未完成采购承诺，计算待采购需求。线性规划模型以“报价+运费-账期资金节省”最小为目标，在满足需求量和供应商可供上限的条件下计算推荐采购组合；再根据提前期反推建议下单日期。",
        "分物料的待采购需求；各供应商推荐采购量、可调整范围、综合成本和备选组合；合同/锁价参考；建议下单日期和到货日期；不可行时的数量或时间缺口。人工确认后进入E采/SAP审批执行。",
        "评分、质量合格率、准时交付率和风险分当前作为业务参考，不直接进入成本目标函数；所有采购和锁价建议必须经人工确认，系统不自动签订合同或生成正式订单。",
        "常规多供应商、多物料采购方案优化计算时间不超过3秒；应能复现需求净算过程，并说明不可行方案的缺口原因。",
        "标书同时要求质量/成本多目标和供应风险量化。项目组需确认是否将质量门槛、供应商集中度或风险上限作为本期模型硬约束，或列入后续阶段。")

    add_scenario_template(doc, marker,
        "成品库存优化与生产补货协同",
        "库存优化场景面向成品，动态回答“应保留多少库存、何时需要安排生产补充、建议补充多少”。系统将月度出货计划、销售订单、实际库存和滚动生产计划统一到月-周-日三级计划链路，并考虑生产补充周期和储罐有效容量。",
        "供应链计划员、生产调度员、仓储管理员、销售运营人员。",
        "每月25日根据下月出货计划发布库存策略；每周分解下一周计划；每日根据实际出入库和未来7天生产计划滚动校验；重大需求、设备或储罐异常时预警。",
        "月度出货计划、销售订单、成品合格/已分配/冻结/待检/质量异常库存、计划生产入库和待发运量；历史出库波动、生产补充周期及波动；成品-储罐映射、工作容量和可用状态。",
        "按成品单独应用动态(R,s,S)策略，按月计算安全库存SS、再订货点s和目标库存S。周计划根据周初可用库存、本周计划流出和已有计划入库计算补货量；日度滚动测算未来7天库存轨迹，识别首个风险日并避免对已纳入排程的计划入库重复补货。",
        "成品SS/s/S参数和发布版本；月度生产补货总量；周生产补货量和入库节奏；日度新增补货建议、风险日和期望可用入库日期；安全库存、再订货点、容量和滞销/呆滞预警。补货建议发送给精益排程模块。",
        "可用库存=合格库存-已分配库存-冻结库存-质量异常库存；计划生产入库只能在预计完工、检验放行并形成可用库存的日期计入；储罐容量按成品允许使用的储罐集合计算，不得直接使用全厂总容量。",
        "给定相同数据和参数时，策略和补货结果可重复；日度计算必须扣除已有计划入库；容量不足、计划延迟或补充周期无法满足时，系统应输出风险日和原因。",
        "标书库存优化范围包含原料、在制品、成品和副产品，现有方案仅对成品形成了可实施设计。需确认本期是否以成品为边界，并在正式蓝图中说明分期计划。")

    add_scenario_template(doc, marker,
        "精益生产排程与异常重排",
        "精益排程将销售需求、库存补货、产能、物料和工艺约束统一到一套计划体系中，以交付可行、安全库存可保障、生产连续和运行效率最大化为原则。现有资料将排程分为未来4-6周粗排和未来24-72小时精排。",
        "供应链计划员、生产调度员、车间计划员、生产负责人。",
        "月度/周度计划发布后生成粗排；每日滚动形成24-72小时精排；销售插单、设备故障、质量事故、物料短缺或计划严重偏差时触发重排。",
        "已发布销售计划与销售订单；成品补货需求和期望入库日期；BOM和工艺路线；产线/设备能力、日历和检修状态；原辅包可用性；储罐/仓储容量；品种切换、清洗、批量和质量约束；订单/客户和产品贡献信息。",
        "粗排按未来4-6周需求和BOM计算日度产量框架和原辅包消耗需求，为采购和供应商发货计划提供依据。精排在粗排框架内，对未来24-72小时订单进行群组化排产，综合记录原料、库存水位、产能、换线、能源时段和贡献信息，形成可执行日计划。",
        "4-6周粗排框架、日原辅包消耗预测、24-72小时车间产线详细计划、生产订单创建建议、订单交付风险、物料/设备/储罐冲突和异常重排方案。发布后将计划传递给MES/SAP执行，实际报工和设备状态回传用于滚动调整。",
        "硬约束必须100%满足；优化目标不得以牺牲设备安全、工艺、质量和物料可用性为代价。方案发布前由业务人员确认；重排应保留原方案、变更原因和影响范围。",
        "对未来72小时、包含数十个订单、多条产线和清洗约束的详细排程，从触发到获得可执行方案的时间应不超过5分钟；输出方案必须满足全部已确认硬约束。",
        "需补充完整产线/设备清单、SKU-产线映射、标准产能和批量、换线清洗矩阵、检修日历、工艺/质量/储罐约束及现行排程样本。在这些资料齐备前，本节仅可作为总体方案，不作为最终约束清单。")

    add_scenario_template(doc, marker,
        "成本统计、分摊与多维盈利分析",
        "成本与盈利分析场景按“数据归集-成本核算-联/副产品分摊-差异分析-边际贡献-多维盈利-管理报告-信号输出”形成闭环，为采购、排程、客户管理和驾驶舱提供可比较的成本与贡献口径。",
        "财务人员、成本核算员、供应链管理人员、管理层。",
        "按日/月核算周期自动运行；成本数据、产量、销量或分摊规则变更时重算；用户可按产品、客户、订单、批次、区域和时间主动分析。",
        "SAP权威的物料价格、成本BOM、费用池、应收/应付、设备折旧和财务过账数据；MES生产、工时和能耗；WMS/TMS仓储和运输数据；销售收入、客户、订单和产品维度。",
        "先将原料、人工、能源、制造、仓储、运输和管理费用归集到统一成本对象，再按经确认的分摊驱动因子分摊联/副产品和共同费用，比较标准成本与实际成本差异，计算产品、订单和客户的边际贡献与盈利。",
        "产品/订单/客户成本与盈利报告；原料、辅料、能源、制造、仓储和运输成本构成；标准-实际成本差异和原因；客户/产品贡献信号，向驾驶舱和排程模块提供决策参考。",
        "所有财务相关数据的权威源必须为中粮SAP；SCOS对财务数据仅做消费、计算和结果回传，不自行替代SAP总账和正式成本过账。每次计算应保留价格期间、分摊规则版本和对账结果。",
        "在给定财务期间、价格版本和分摊规则时，成本结果应可重复和可与SAP对账；差异必须能钻取到费用池、业务量和分摊驱动因子。",
        "需财务部门确认成本对象、费用池、标准/实际成本口径、联/副产品分摊规则、边际贡献公式和SAP回传/对账方式。")

    add_scenario_template(doc, marker,
        "供应链全程追溯与质量协同",
        "全程追溯场景以批次为主线，连接供应商批次、原料入库、投料、生产、检验、成品入库、销售发运和客户交付，支持从原料到客户的正向追溯和从客户/成品到原料的反向追溯。",
        "质量、生产、仓储、物流、销售和供应链管理人员。",
        "用户输入原料批次、生产批次、成品批次、销售订单、发运单或客户后触发查询；LIMS/QMS质量异常事件发生时触发联动。",
        "MES投料、生产和报工批次；WMS原料/成品入出库与库存批次；LIMS/QMS检验结果和质量状态；SAP采购、生产和销售单据；TMS发运和交付信息。",
        "系统以统一批次ID和单据号建立“供应批-投料批-生产批-成品批-发运批”关联网络，并将检验、放行、隔离、返工和报废状态附着到批次关系上。质量异常发生时，根据批次影响范围查找在制、库存和已发运对象。",
        "正向/反向追溯树、关联单据和检验结果；影响成品、库存、订单、客户和供应商清单；隔离、召回、放行或处置建议；向驾驶舱推送重大质量预警。",
        "追溯关系不得仅依赖产品名称和日期模糊匹配；批次编码、上下游关联和质量状态更正必须有权限、原因和审计记录。追溯系统只提供影响分析和处置协同，正式质量放行/召回决定仍由责任部门执行。",
        "对预置批次数据集进行正向和反向追溯时，返回的批次、单据和客户范围应完整、无重复和无错连；质量异常应能联动到预警与处置流程。",
        "需确认原料/生产/成品批次编码规则、MES/WMS/LIMS/SAP/TMS关联键、质量隔离/放行/召回流程及追溯时效指标。")

    add_heading(doc, marker, "智能化场景验收与业务边界", 2)
    add_table(doc, marker, ["验收维度", "蓝图初稿口径", "验收前必须补充"], [
        ("正确性", "净需求、库存位置、成本、批次和指标口径可与手工/权威系统结果对账", "基准数据集、预期结果和允许偏差"),
        ("可行性", "排程硬约束100%满足；采购/库存不可行时输出缺口原因", "完整硬/软约束清单和优先级"),
        ("性能", "采购常规计算≤3秒；72小时排程≤5分钟；驾驶舱加载≤3秒、钻取≤2秒", "数据规模、并发条件、测试环境和硬件配置"),
        ("可解释/审计", "保存输入快照、参数、模型版本、方案、人工调整与发布记录", "日志保留期、调整原因码和审批流"),
        ("业务效果", "持续跟踪交付、库存、缺货、采购成本、计划稳定性和预警闭环", "上线前基线、目标值、统计周期和指标责任人"),
    ], [1200, 3900, 3180], font=8.2)
    add_note(doc, marker, "本章已将现有资料尽可能转化为可评审正文。库存品类范围、采购多目标/风险、排程约束、成本分摊和批次编码是正式定稿前的高优先级评审项。", label="编制说明")


def write_chapter_4(doc, marker):
    add_heading(doc, marker, "SCOS技术蓝图交付物", 1, page_break=True)
    add_heading(doc, marker, "SCOS总体技术架构", 2)
    add_p(doc, marker, "SCOS总体架构遵循“数据统一供给、业务模块解耦、模型服务化、结果统一发布、全程可审计”原则。SAP、MES、WMS、TMS、QMS/LIMS、EAM/EMS、E采和外部期货价格服务是业务事实来源；统一数据服务层将主数据、计划、交易、库存、价格、成本、批次和设备数据形成标准数据服务；SCOS各模块基于这些数据开展计算和协同。")
    if ARCH_IMG.exists():
        add_picture(doc, marker, ARCH_IMG, "图4-1  SCOS整体架构（来源：现有《SCOS整体架构》资料，正式稿需统一图号与版本）")
    add_table(doc, marker, ["架构层", "主要组件", "技术责任"], [
        ("数据源层", "SAP、MES、WMS、TMS、QMS/LIMS、EAM/EMS、E采、CRM、外部行情服务", "维护主数据、单据、批次、设备、库存、财务和执行事实。"),
        ("集成与数据层", "统一数据服务层、实时数据平台、企业服务总线、API网关/治理中心", "标准化数据包、API和事件，统一认证、限流、日志和服务治理。"),
        ("SCOS业务层", "采购优化、成品库存、精益排程、成本分析、追溯协同、驾驶舱", "管理业务参数、方案、审批、发布和执行反馈，各模块独立部署和迭代。"),
        ("模型与计算层", "采购LP、库存(R,s,S)、排程求解、KPI计算、情景编排", "实现标准化输入输出、任务调度、超时、可行性校验、版本和回退。"),
        ("应用层", "业务工作台、方案对比、预警中心、驾驶舱、沙盘推演、移动端", "面向不同角色提供查询、计算、调整、审批和报告能力。"),
        ("运维治理层", "统一身份权限、配置中心、日志、链路追踪、监控告警、备份容灾", "保障系统安全、可用、可观测、可恢复和可审计。"),
    ], [1300, 3100, 3880], font=8.1)
    add_heading(doc, marker, "架构原则与系统边界", 3)
    add_table(doc, marker, ["原则", "蓝图表述"], [
        ("消费而非采集", "SCOS通过统一数据服务层和实时数据平台获取数据，自身不重复建设与统一平台重叠的采集、清洗和存储管道。"),
        ("直连例外管理", "确需与SAP等系统进行深度、双向业务集成时必须单独论证，按项目标准协议实现，并在API治理中心注册。"),
        ("高内聚、松耦合", "各业务模块可作为独立服务部署和迭代，不共享私有数据表，通过统一数据契约、API或事件进行交互。"),
        ("人在回路", "模型输出是建议方案，与正式业务单据隔离；经业务确认和审批后才能发布或写入执行系统。"),
        ("权威源与可审计", "财务价格、费用分摊等数据以SAP为权威源；每次模型计算和数据发布均保留版本和审计记录。"),
    ], [1700, 6580], font=8.5)

    add_heading(doc, marker, "SCOS数据流架构", 2)
    add_p(doc, marker, "SCOS数据流以“源系统数据-统一数据底稿-业务/模型计算-人工审批-结果发布-执行反馈”为主线。《SCOS逻辑》资料中已初步明确了SAP、SCOS、TMS、MES和WMS之间的业务关系：SAP提供销售计划、BOM、供应商、合同、价格和财务数据；TMS提供在途和发货数据；MES提供生产执行；WMS提供成品库存事实。SCOS在统一底稿上完成采购、排程、库存、成本和驾驶舱处理。")
    add_table(doc, marker, ["数据阶段", "处理内容", "质量/控制要求"], [
        ("获取/订阅", "通过统一数据服务API获取主数据和交易数据；通过实时平台订阅产量、库存、液位、产线和位置数据。", "记录数据源、时间戳、版本和接口调用状态。"),
        ("标准化/快照", "将不同系统的物料、客户、供应商、单位、组织、状态和时间口径映射为SCOS标准数据包。", "对编码、必填项、单位、重复、参照关系和时效进行校验。"),
        ("计算/模拟", "按发布数据版本生成计算任务，调用相应规则或模型，保存输入快照和模型参数。", "输入不完整、数据过期或模型不可行时终止并返回原因。"),
        ("审批/发布", "用户对方案进行对比和调整，按权限和审批流程形成已发布版本。", "保留人工调整前后差异、原因、审批人、时间和发布对象。"),
        ("执行/反馈", "已发布计划或建议通过企业服务总线/API发送到SAP、MES、E采等执行系统，执行状态回传SCOS。", "保障幂等、重试、超时、失败补偿和业务对账。"),
    ], [1300, 4650, 2330], font=8.1)
    add_note(doc, marker, "《交互数据初稿》中存在“SCOS审批后直接创建SAP生产订单”和“采购订单无需与SAP对接”等初稿口径。这些涉及系统交易边界和主数据一致性，正式蓝图必须由SAP/E采/MES系统Owner专项确认，本稿不将其写成无条件承诺。")

    add_heading(doc, marker, "SCOS数据域与核心实体", 2)
    add_p(doc, marker, "SCOS不以建设新的企业主数据系统为目标，但必须在模型计算与方案发布过程中维护稳定的业务键、数据版本和引用关系。下表为现阶段可从《SCOS整体架构》《SCOS逻辑》和《交互数据初稿》中沉淀的核心数据域。")
    add_table(doc, marker, ["数据域", "核心实体", "关键属性/版本", "候选权威源"], [
        ("主数据", "物料、供应商、客户、仓库、储罐、生产设备、产线", "编码、名称、分类、组织、单位、状态、生效/失效时间", "SAP/MES/WMS/主数据平台（待确认）"),
        ("需求与销售", "销售预测/业计划、月度出货计划、销售订单、交付承诺", "计划版本、发布状态、客户、SKU、数量、需求/承诺日期", "SAP/CRM/计划系统"),
        ("生产与排程", "粗排、车间详细排程、生产订单、报工、设备日历", "方案版本、产线/设备、SKU、数量、开完工、状态、来源需求", "SCOS/MES/SAP"),
        ("库存与容量", "库存快照、库存流水、批次、库存分配、储罐、计划入出库、库存策略", "合格/已分配/冻结/待检/异常、数量、时点、有效容量、SS/s/S版本", "WMS/MES/LIMS/实时平台/SCOS"),
        ("采购与供应", "原材料需求、采购合同/订单、报价、供货能力、账期、发货/在途/到货、供应商考核", "物料、供应商、周期、未完成量、报价有效期、提前期、运费、MOQ、准时率、质量和风险", "SAP/E采/TMS/LIMS/QMS/外部行情"),
        ("财务与成本", "原料价格、成本BOM、费用池、资金成本、应收/应付、设备折旧、成本结果", "币种、期间、价格版本、成本中心、分摊规则、过账/对账状态", "SAP（权威源）"),
        ("质量与追溯", "供应批、投料批、生产批、成品批、检验结果、发运批、质量事件", "批次ID、上下游批次、检验/放行/隔离状态、时间、单据号", "MES/WMS/LIMS/QMS/SAP/TMS"),
        ("模型与方案", "计算任务、输入快照、参数集、模型版本、结果方案、人工调整、发布记录", "任务ID、触发人/时间、数据版本、模型版本、状态、可行性、耗时、发布对象", "SCOS"),
    ], [1000, 2550, 2900, 1830], font=7.8)
    add_heading(doc, marker, "SCOS与SAP主要交互数据底稿", 3)
    add_table(doc, marker, ["方向", "数据类型", "当前说明"], [
        ("SAP→SCOS", "物料、供应商、客户、仓库、储罐、生产设备主数据", "作为SCOS计算和方案对象的基础参照数据。"),
        ("SAP→SCOS", "历史出货、原料价格、未关闭合同、未清销售订单、未结采购/生产订单", "用于需求、采购、排程和库存算法的交易事实。"),
        ("SAP→SCOS", "物料BOM、制造BOM、成本BOM、费用池、应收应付、设备折旧", "用于物料净算、工艺/工时、成本核算和资金指标。"),
        ("SCOS→SAP", "已审批生产计划/生产订单建议", "是否由SCOS直接创建SAP生产订单需专项确认。"),
        ("SCOS→SAP", "成本分摊结果", "回传订单行成本分摊结果的方式、频率和对账机制待确认。"),
        ("其他系统→SCOS", "销售预测、供应商准时率/质量/产能/MOQ、成品在途、库存持有成本、资金成本", "《交互数据初稿》尚未明确具体来源系统和“有无”状态，需调研确认。"),
    ], [1250, 3600, 3430], font=8.0)

    add_heading(doc, marker, "SCOS模型与计算服务", 2)
    add_p(doc, marker, "为保证各业务模块可独立部署、可测试和可回退，优化算法应封装为标准模型服务。每个服务接收明确的数据版本、参数集和触发信息，返回任务状态、结果方案、可行性、性能耗时和解释信息。")
    add_table(doc, marker, ["服务", "输入摘要", "处理/目标", "输出摘要", "性能/状态"], [
        ("采购优化", "待采购需求、供应商报价/可供量/运费/账期、提前期", "线性规划；最小化报价+运费-账期节省", "推荐/备选组合、综合成本、下单/到货建议、缺口", "常规计算≤3秒；风险约束待确认"),
        ("成品库存", "需求、可用库存、计划入库、补充周期、储罐容量", "动态(R,s,S)；月-周-日协同与7天滚动", "SS/s/S、补货量、风险日、期望入库日期、容量风险", "每日计算；策略原则上按月发布"),
        ("精益排程", "需求/补货、BOM/工艺、产能、物料、储罐、设备日历、切换/质量约束", "粗排+精排；满足交付、库存、连续生产和效率目标", "4-6周框架、24-72小时详细计划、冲突、风险和重排方案", "72小时≤5分钟；硬约束100%"),
        ("KPI与健康度", "业务数据快照、指标公式、阈值和预警规则", "指标计算、红黄绿映射和分级预警", "KPI、健康度、预警、钻取维度和快照时间", "运营数据延迟≤5秒；分析数据≤1小时"),
        ("沙盘编排", "基线快照、情景参数、被调模型服务清单", "在沙盘隔离环境调用多个模型并汇总影响", "当前-模拟对比、不可行原因、可采纳建议", "继承被调模型SLA；编排超时待确认"),
    ], [1050, 2000, 2300, 1950, 980], font=7.5)
    add_heading(doc, marker, "模型运行与降级机制", 3)
    add_p(doc, marker, "模型计算任务应包含“待执行、执行中、成功、不可行、数据异常、超时、取消”等状态。当输入缺失或数据时效不满足时，系统不启动模型；当输出不可行时，应返回约束冲突或缺口；当求解器故障或性能连续不达标时，应允许回退到上一已验证模型版本，并保留切换记录。")
    add_note(doc, marker, "生产排程求解器尚未完成选型。正式蓝图需说明主求解器、商业授权、硬件资源、备用求解路径和故障替换机制。")

    add_heading(doc, marker, "数据发布、版本与审计规范", 2)
    add_p(doc, marker, "SCOS中所有能够影响业务执行的计划、策略和参数均应按版本管理。草稿版本用于计算和人工调整，审批中版本不允许直接修改，已发布版本作为下游执行和数据消费的唯一正式来源，已失效版本仅供追溯。")
    add_table(doc, marker, ["对象", "必须记录的版本信息", "发布/变更要求"], [
        ("需求/计划", "计划类型、期间、版本号、数据来源、生成时间、发布状态", "变更后评估对采购、库存和排程的影响，仅调整未执行部分。"),
        ("模型参数", "参数集版本、适用对象、生效期、修改人、修改原因、审批人", "SS/s/S、阈值、分摊规则等参数修改后不追溯篡改已发布历史结果。"),
        ("模型版本", "模型名称、版本、代码/求解器版本、发布包、验证数据集、性能结果", "新版本上线前必须通过回归和性能测试，并支持快速回退。"),
        ("方案与结果", "计算任务ID、输入快照、参数/模型版本、原始结果、人工调整、审批/发布记录", "方案发布后保留不可篡改快照；撤回/失效通过新状态和新版本实现。"),
        ("接口数据", "接口/事件版本、消息ID、业务键、发送/接收时间、处理状态、重试次数", "涉及业务写入的接口必须幂等，失败时支持重试、补偿和人工对账。"),
    ], [1300, 4000, 2980], font=8.0)

    add_heading(doc, marker, "安全、运维与非功能要求", 2)
    add_table(doc, marker, ["类别", "蓝图要求", "待设计/待确认"], [
        ("身份与权限", "接入企业统一身份认证，按角色和数据范围授权；计算、调整、审批、发布、回退和配置权限分离。", "统一身份产品、角色-权限矩阵、单点登录和账号同步机制。"),
        ("数据与接口安全", "按数据分级分类控制查看、导出和发布；API使用认证、加密传输、白名单、限流和审计日志。", "等保/中粮安全基线、密钥管理、数据脱敏和导出审批要求。"),
        ("可用性与容灾", "各模块可独立部署、扩容和回退；关键服务实施健康检查、故障转移、备份恢复和容灾演练。", "部署拓扑、节点数、RTO/RPO、备份周期和容灾环境。"),
        ("可观测性", "采集应用、接口、模型任务、数据质量、消息积压、资源和业务指标，实现日志-指标-链路统一定位。", "监控平台、指标阈值、告警升级、日志保留期和运维工单集成。"),
        ("性能与并发", "采购优化≤3秒；72小时排程≤5分钟；主视图加载≤3秒、钻取≤2秒；至少支持50并发用户。", "为每项指标固定数据规模、测试环境、网络、缓存、并发和统计口径。"),
        ("质保与服务", "自正式上线验收合格起提供至少12个月免费质保；建立故障分级、响应和升级机制。", "完整SLA、恢复时间、服务窗口、联系矩阵、运维工具和备件管理。"),
    ], [1200, 4350, 2730], font=8.0)

    add_heading(doc, marker, "技术蓝图待确认项", 2)
    add_table(doc, marker, ["优先级", "待确认项", "影响", "建议责任方"], [
        ("P0", "统一数据服务层、实时数据平台、ESB和API网关的实际产品与能力边界", "无法固化总体架构、接口责任和工期", "业主IT+集成厂商+SCOS"),
        ("P0", "各源系统字段级接口清单、权威数据源和双向写入范围", "数据实体映射和业务闭环无法设计/验收", "系统Owner+业务Owner"),
        ("P0", "排程完整约束、求解器选型、授权和备用求解路径", "无法完成5分钟/100%可行性设计和验收", "算法组+生产工艺+IT"),
        ("P0", "库存优化品类范围、采购风险/多目标与标书偏差处理决议", "蓝图范围和验收范围可能冲突", "项目组+业主"),
        ("P1", "成本分摊、批次编码和质量处置流程", "成本和追溯只能保留在总体方案层", "财务+质量+生产+仓储"),
        ("P1", "部署资源、网络区域、身份安全、日志、备份容灾和RTO/RPO标准", "无法完成部署图、容量和安全设计", "业主IT+信息安全+供应商"),
    ], [650, 3650, 2400, 1580], font=7.8)
    add_note(doc, marker, "本章已达到可进行业务与技术联合评审的初稿粒度。字段级接口、实体主键、部署拓扑、安全基线和求解器选型应在专项方案或后续版本中继续细化。", label="编制说明")


def write_chapters_5_to_8(doc, marker):
    add_heading(doc, marker, "SCOS集成接口规范", 1, page_break=True)
    add_p(doc, marker, "SCOS对外集成遵循统一数据服务、企业服务总线和API治理规范。下表是根据现有《交互数据初稿》《SCOS整体架构》及各模块方案形成的业务级接口清单。接口编号、字段、协议、路由、频率和安全配置将在详细设计阶段与各系统Owner联合确认。")
    add_heading(doc, marker, "数据获取与订阅接口", 2)
    add_table(doc, marker, ["对接方", "接口/数据包", "主要内容", "建议方式", "频率/触发", "状态"], [
        ("SAP", "主数据", "物料、供应商、客户、仓库、储罐、生产设备", "统一数据服务API", "日/变更触发", "待确认"),
        ("SAP", "销售与交易数据", "销售计划、历史出货、未清销售订单、产品单价", "API/定时数据服务", "日/变更触发", "待确认"),
        ("SAP", "BOM与生产数据", "物料BOM、制造BOM、未结生产订单", "API/定时数据服务", "日/版本变更", "待确认"),
        ("SAP", "采购与价格数据", "原料价格、未关闭合同、未结采购订单", "API/定时数据服务", "小时/日/变更", "待确认"),
        ("SAP", "财务与成本数据", "成本BOM、费用池、应收应付、设备折旧、资金成本", "API/数据服务", "小时/月结", "待确认"),
        ("MES", "生产执行与设备状态", "生产订单、报工、投料/产出批次、产线/设备状态", "事件+查询API/实时平台", "实时/秒级", "待确认"),
        ("WMS", "库存与出入库", "成品/原料库存、分配/冻结/待检状态、批次、出入库", "事件+查询API/实时平台", "实时/日对账", "待确认"),
        ("TMS", "运输与在途", "原料/成品在途、发货、车辆位置、预计到达时间和运费", "事件+查询API", "实时/状态变更", "待确认"),
        ("LIMS/QMS", "质量与检验", "检验结果、放行/隔离状态、供应商合格率、质量事件", "事件+查询API", "检验/状态变更", "待确认"),
        ("E采", "采购执行数据", "采购审批、订单、供应商发货和执行状态", "事件+查询API", "状态变更", "待确认"),
        ("CRM", "客户需求事件", "客户需求、订单变更、客户贡献信息", "ESB事件+GET/POST", "事件驱动", "待确认"),
        ("外部行情", "期货/大宗商品价格", "当前价格、品种、合约、单位、币种、时间戳", "安全API", "小时/日", "数据源待确认"),
    ], [850, 1500, 2750, 1450, 950, 780], font=7.2)
    add_heading(doc, marker, "结果下发与闭环接口", 2)
    add_table(doc, marker, ["接收方", "下发对象", "触发条件", "闭环回传", "状态"], [
        ("MES/SAP", "已审批生产计划/生产订单建议", "排程方案发布", "订单号、接收状态、生产报工和计划偏差", "交易边界待确认"),
        ("E采/SAP", "已审批采购建议、建议下单/到货日期", "采购方案发布", "采购订单、发货、在途、到货和完成状态", "执行系统待确认"),
        ("SAP", "成本分摊结果", "成本计算审核通过", "过账/对账状态及差异", "回传方式待确认"),
        ("驾驶舱", "采购、库存、排程、成本、追溯方案和预警", "方案/预警生成或状态变更", "钻取、处置和闭环状态", "内部标准服务"),
        ("数字孪生", "空间对象业务态势、预警和钻取链接", "实时状态/预警变更", "用户交互或选中对象回调", "范围待确认"),
    ], [1000, 2550, 1700, 2200, 830], font=7.8)
    add_heading(doc, marker, "接口通用技术要求", 2)
    add_table(doc, marker, ["项目", "通用要求"], [
        ("版本与契约", "每个API/事件应有唯一编号、版本、请求/响应Schema、必填项、枚举、错误码和退役计划。"),
        ("安全", "通过统一认证、加密传输、网络访问控制、白名单、限流、请求签名和审计日志保护接口。"),
        ("幂等与一致性", "写入接口必须通过消息ID+业务键保证幂等；接口失败不得造成部分重复或无状态的中间数据。"),
        ("重试和补偿", "短暂故障按退避策略重试；超过上限后进入异常队列/工单，由系统或人工对账补偿。"),
        ("时效与监控", "对调用量、成功率、响应时间、超时、重试、数据时效和消息积压建立监控与告警。"),
        ("对账", "对库存快照与流水、订单状态、成本结果和计划发布等关键接口建立定期对账和差异处理。"),
    ], [1600, 6680], font=8.4)

    add_heading(doc, marker, "SCOS业务场景-智能化覆盖率映射表", 1, page_break=True)
    add_heading(doc, marker, "智能化场景清单", 2)
    add_table(doc, marker, ["序号", "场景名称", "触发方式", "智能技术", "输入来源", "输出去向", "纳入验收"], [
        ("1", "全局态势监测与健康度评估", "定时/实时", "指标计算+规则", "统一数据服务/实时平台", "驾驶舱", "是"),
        ("2", "分级预警、钻取溯源与处置闭环", "事件/规则", "规则+关联分析", "各模块预警/业务明细", "预警中心/工单", "是"),
        ("3", "需求/供应/成本沙盘推演", "人工", "情景编排+模型调用", "基线快照/情景参数", "驾驶舱/各优化模块", "是"),
        ("4", "原材料待采购需求净算", "月/周/人工", "净需求规则", "原材料需求/未完成采购承诺", "采购优化模块", "是"),
        ("5", "多供应商采购组合优化", "月/周/人工", "线性规划", "待采购需求/供应条件", "采购审批/E采/SAP", "是"),
        ("6", "成品安全库存与生产补货", "月/周/日", "动态(R,s,S)", "需求/库存/生产计划/储罐", "排程/驾驶舱", "是"),
        ("7", "4-6周粗排与24-72小时精排", "周/日/异常", "约束优化", "需求/BOM/产能/物料/设备", "MES/SAP/采购/库存", "待约束齐备后纳入"),
        ("8", "成本分摊与多维盈利分析", "日/月/人工", "规则+多维分析", "SAP/MES/WMS/TMS", "SAP/驾驶舱/排程", "待口径齐备后纳入"),
        ("9", "批次正反向追溯与质量联动", "人工/质量事件", "图关联+规则", "MES/WMS/LIMS/SAP/TMS", "追溯/预警/处置闭环", "待批次规则齐备后纳入"),
    ], [500, 2200, 950, 1250, 1500, 1250, 630], font=7.2)
    add_heading(doc, marker, "场景验收指标", 2)
    add_table(doc, marker, ["场景", "验收指标", "初稿目标", "验证方式"], [
        ("驾驶舱", "主视图加载/钻取响应", "≤3秒/≤2秒", "在约定数据规模和50并发条件下压测"),
        ("驾驶舱", "关键运营/成本资金指标时效", "≤5秒/≤1小时", "比较源系统时间戳与驾驶舱可见时间"),
        ("采购优化", "常规方案计算时间", "≤3秒", "使用典型多供应商、多物料数据连续测试"),
        ("采购优化", "净需求与约束正确性", "与基准结果一致", "手工核算待采购量，检查推荐量和可供上限"),
        ("成品库存", "计划入库防重、容量和风险日正确性", "与预期结果一致", "使用已知需求、库存、计划入库和储罐场景回放"),
        ("精益排程", "72小时详细排程时间/硬约束", "≤5分钟/100%", "约束齐备后用典型订单、多产线和清洗场景验证"),
        ("成本/追溯", "对账或批次链完整性", "口径齐备后确定", "与SAP成本基准对账；用预置批次链正反向查询"),
    ], [1500, 2700, 1500, 2580], font=8.0)
    add_heading(doc, marker, "智能化场景覆盖率", 2)
    add_table(doc, marker, ["统计项", "数量", "说明"], [
        ("蓝图已定义智能化场景", "9", "覆盖驾驶舱、沙盘、采购、库存、排程、成本和追溯。"),
        ("当前可直接纳入验收场景", "6", "驾驶舱2项、沙盘1项、采购2项、成品库存1项；仍需确认基准数据。"),
        ("待资料齐备后纳入验收场景", "3", "精益排程、成本分析和批次追溯。"),
        ("最终智能化覆盖率", "待确认", "应在本期范围、验收口径和数据条件确认后计算，本稿不提前承诺比例。"),
    ], [2600, 1000, 4680], font=8.4)

    add_heading(doc, marker, "SCOS非功能性需求响应表", 1, page_break=True)
    add_table(doc, marker, ["类别", "技术要求/目标", "SCOS响应方案", "状态"], [
        ("架构合规", "依赖统一数据服务层和实时数据平台", "各模块不重复建设数据采集/清洗管道；直连例外单独论证并注册治理。", "满足，平台产品待确认"),
        ("模块化", "模块高内聚、松耦合，可独立部署迭代", "采购、库存、排程、成本、追溯和驾驶舱通过API/事件交互。", "满足"),
        ("算法性能", "采购常规计算≤3秒；72小时排程≤5分钟", "建立异步任务、超时、性能基准、模型版本和回退机制。", "方案已响应，基准待确认"),
        ("结果可行性", "排程硬约束100%满足", "求解后执行独立可行性检查；不可行时返回约束冲突，不发布计划。", "约束清单待确认"),
        ("驾驶舱性能", "主视图≤3秒、钻取≤2秒、至少50并发", "按指标层级、时间粒度和角色视图设计缓存和查询，并建立压测基准。", "满足，测试条件待确认"),
        ("数据时效", "关键运营数据延迟≤5秒；成本/资金≤1小时", "实时数据通过消息订阅，成本/资金按数据服务定时获取，记录源时间戳。", "满足，接口SLA待确认"),
        ("可配置性", "指标、阈值、预警和看板可配置", "配置对象统一管理版本、生效期、权限、审批和回退。", "满足"),
        ("安全与审计", "统一身份、最小权限、接口安全、全程留痕", "实现分角色/分数据范围授权，审计计算、调整、审批、发布、导出和接口操作。", "原则满足，安全基线待确认"),
        ("可用性/容灾", "模块故障隔离、快速回退、备份恢复和求解器应急", "实施健康检查、故障转移、数据备份、容灾演练和模型/求解器替代路径。", "RTO/RPO和拓扑待确认"),
        ("运维与质保", "正式验收后至少12个月免费质保", "建立P1/P2等故障分级、响应、升级、恢复和服务报告机制。", "满足，完整SLA待确认"),
    ], [1100, 2500, 3700, 980], font=7.8)

    add_heading(doc, marker, "SCOS蓝图交付物清单", 1, page_break=True)
    add_p(doc, marker, "下表列出本轮蓝图设计建议形成的主要交付物。“本稿已覆盖”表示已有结构和主要内容，不表示已完成业务确认或项目验收。")
    add_table(doc, marker, ["序号", "交付物", "主要内容", "当前状态", "后续动作"], [
        ("1", "SCOS业务流程总览", "端到端业务总图及关键流程", "现有第二章", "业务联合评审"),
        ("2", "关键业务流程说明", "年/月/日计划、需求、排程、原辅包和报工流程", "现有第二章", "补异常和跨系统流程"),
        ("3", "系统边界、功能架构和权限", "系统分工、功能模块、角色和业务规则", "现有第二章", "完成范围决议后修订"),
        ("4", "智能化应用场景说明书", "驾驶舱、沙盘、采购、库存、排程、成本和追溯场景", "本稿已覆盖", "关闭第三章待确认项"),
        ("5", "SCOS总体技术架构", "分层架构、责任边界和架构原则", "本稿已覆盖", "确认业主平台产品与拓扑"),
        ("6", "SCOS数据流与数据域设计", "数据主线、数据域、核心实体和SAP交互底稿", "本稿已覆盖", "补字段字典和权威源矩阵"),
        ("7", "模型与计算服务设计", "采购、库存、排程、KPI和沙盘服务", "本稿已覆盖", "补排程约束、求解器和JSON Schema"),
        ("8", "数据发布、版本与审计规范", "计划、参数、模型、方案和接口版本", "本稿已覆盖", "与质量与安全标准对齐"),
        ("9", "SCOS集成接口清单", "获取/订阅、结果下发和接口通用规范", "业务级已覆盖", "与各系统Owner补字段级契约"),
        ("10", "智能化场景及验收映射", "9个场景、验收指标和覆盖率口径", "本稿已覆盖", "确认范围、基线和最终比例"),
        ("11", "非功能需求响应表", "架构、性能、时效、安全、容灾、运维和质保", "原则级已覆盖", "补环境、容量、RTO/RPO和SLA"),
        ("12", "待确认事项与责任矩阵", "业务范围、数据、接口、算法、安全和部署缺口", "本稿分散列示", "评审后汇总成闭环台账"),
    ], [550, 2100, 2900, 1300, 1430], font=7.7)
    add_note(doc, marker, "建议下一步以本版本为联合评审底稿，按“范围决议-业务规则-数据与接口-算法与验收-部署与安全”顺序关闭P0/P1事项，再发布正式蓝图。", label="后续建议")


def set_update_fields(doc):
    settings = doc.settings._element
    node = settings.find(qn("w:updateFields"))
    if node is None:
        node = OxmlElement("w:updateFields"); settings.append(node)
    node.set(qn("w:val"), "true")


def restore_template_parts():
    fd, temp_name = tempfile.mkstemp(
        prefix="scos_v02_", suffix=".docx", dir=OUTPUT.parent
    )
    os.close(fd)
    temp_path = Path(temp_name)
    try:
        with ZipFile(SOURCE, "r") as source_zip, ZipFile(OUTPUT, "r") as output_zip:
            source_names = set(source_zip.namelist())
            with ZipFile(temp_path, "w") as result_zip:
                for info in output_zip.infolist():
                    if info.filename in PRESERVE_TEMPLATE_PARTS and info.filename in source_names:
                        data = source_zip.read(info.filename)
                    else:
                        data = output_zip.read(info.filename)
                    result_zip.writestr(info, data)
        os.replace(temp_path, OUTPUT)
    finally:
        if temp_path.exists():
            temp_path.unlink()


def main():
    digest = hashlib.sha256(SOURCE.read_bytes()).hexdigest()
    if digest != EXPECTED_SHA:
        raise SystemExit(f"Reference changed: {digest}")
    doc = Document(SOURCE)
    update_visible_toc(doc)
    marker = remove_between(doc, "SCOS驾驶舱场景推演功能应用场景说明书", "SCOS集成接口规范")
    marker = remove_from_to_sectpr(doc, marker)
    write_chapter_3(doc, marker)
    write_chapter_4(doc, marker)
    write_chapters_5_to_8(doc, marker)
    set_update_fields(doc)
    doc.core_properties.title = "SCOS中粮太仓项目蓝图方案v0.2-业务蓝图初稿"
    doc.core_properties.subject = "SCOS智能化应用场景与技术蓝图"
    doc.save(OUTPUT)
    restore_template_parts()
    print(OUTPUT)


if __name__ == "__main__":
    main()
